"""
Script de révision de documents Word avec conservation du formatage.
Supporte la correction orthographique, la traduction et d'autres commandes via OpenAI.
"""

import os
import sys
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
import json
from dotenv import load_dotenv
from langdetect import detect, DetectorFactory
from datetime import datetime
import difflib

# Configurer l'encodage UTF-8 pour Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

# Charger le fichier .env s'il existe
load_dotenv()

# Fixer le seed pour langdetect (pour avoir des résultats reproductibles)
DetectorFactory.seed = 0


class DocumentReviewer:
    """Classe pour réviser des documents Word avec conservation du formatage."""
    
    def __init__(self, api_key: Optional[str] = None, model: Optional[str] = None):
        """
        Initialise le reviewer avec une clé API OpenAI.
        
        Args:
            api_key: Clé API OpenAI (si None, utilise la variable d'environnement OPENAI_API_KEY)
            model: Modèle OpenAI à utiliser (si None, utilise OPENAI_MODEL depuis .env ou gpt-4o par défaut)
        """
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError(
                "Clé API OpenAI requise. Définissez OPENAI_API_KEY ou passez-la en paramètre."
            )
        
        # Charger le modèle depuis .env ou utiliser le paramètre ou la valeur par défaut
        self.model = model or os.getenv("OPENAI_MODEL") or "gpt-4o"
        
        self.client = OpenAI(api_key=self.api_key)
        self.conversation_history: List[Dict] = []
        self.current_document: Optional[Document] = None
        self.current_path: Optional[Path] = None
        self.paragraphs_cache: List[str] = []
        self.detected_language: Optional[str] = None
        self.log_file: Optional[Path] = None
        self.current_instruction: Optional[str] = None
        self.paragraphs_with_images: List[int] = []  # Liste des numéros de paragraphes avec images
        self.initial_image_count: int = 0
        
    def load_document(self, file_path: str) -> None:
        """
        Charge un document Word en mémoire.
        
        Args:
            file_path: Chemin vers le fichier .doc ou .docx
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Le fichier {file_path} n'existe pas.")
        
        if path.suffix.lower() not in ['.doc', '.docx']:
            raise ValueError("Le fichier doit être au format .doc ou .docx")
        
        self.current_document = Document(file_path)
        self.current_path = path
        self.paragraphs_cache = [p.text for p in self.current_document.paragraphs]
        
        # Détecter la langue du document
        self._detect_language()
        
        print(f"✓ Document chargé: {path.name}")
        print(f"  Nombre de paragraphes: {len(self.paragraphs_cache)}")
        print(f"  Modèle OpenAI: {self.model}")
        if self.detected_language:
            print(f"  Langue détectée: {self._get_language_name(self.detected_language)}")
        
        # Initialiser le fichier de log
        self._init_log_file()
        
        # Compter les images initiales
        self._count_images()
        
    def _count_images(self) -> None:
        """
        Compte toutes les images dans le document et identifie les paragraphes qui en contiennent.
        """
        self.initial_image_count = 0
        self.paragraphs_with_images = []
        
        for i, paragraph in enumerate(self.current_document.paragraphs):
            if self._has_images(paragraph):
                self.paragraphs_with_images.append(i + 1)
                # Compter le nombre d'images dans ce paragraphe
                for run in paragraph.runs:
                    if hasattr(run, '_element'):
                        for child in run._element:
                            if 'drawing' in child.tag or 'pict' in child.tag:
                                self.initial_image_count += 1
        
        if self.initial_image_count > 0:
            print(f"  Images trouvées: {self.initial_image_count} image(s) dans {len(self.paragraphs_with_images)} paragraphe(s)")
            print(f"  ⚠️  Les paragraphes avec images ne seront PAS modifiés pour les préserver")
    
    def _has_images(self, paragraph) -> bool:
        """
        Vérifie si un paragraphe contient des images.
        
        Args:
            paragraph: Paragraphe docx à vérifier
            
        Returns:
            True si le paragraphe contient des images
        """
        try:
            for run in paragraph.runs:
                # Vérifier les inline shapes (images)
                if hasattr(run, '_element'):
                    for child in run._element:
                        # drawing est le tag pour les images/formes inline
                        if 'drawing' in child.tag or 'pict' in child.tag:
                            return True
        except:
            pass
        return False
    
    def _extract_styles_map(self, paragraph) -> List[Dict]:
        """
        Extrait une carte détaillée des styles du paragraphe avec positions exactes.
        
        Args:
            paragraph: Paragraphe docx
            
        Returns:
            Liste de dictionnaires avec le style et la plage de caractères
        """
        styles_map = []
        char_position = 0
        
        for run in paragraph.runs:
            if run.text:  # Inclure même les runs vides pour les images
                style = {
                    'start': char_position,
                    'end': char_position + len(run.text),
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb if run.font.color.rgb else None,
                }
                styles_map.append(style)
                char_position += len(run.text)
        
        return styles_map
    
    def _map_styles_to_new_text(self, original_text: str, new_text: str, styles_map: List[Dict]) -> List[Dict]:
        """
        Mappe intelligemment les styles de l'ancien texte vers le nouveau en utilisant difflib.
        
        Args:
            original_text: Texte original
            new_text: Nouveau texte
            styles_map: Carte des styles de l'original
            
        Returns:
            Nouvelle carte de styles adaptée au nouveau texte
        """
        if not styles_map:
            return []
        
        # Utiliser SequenceMatcher pour comprendre les changements
        matcher = difflib.SequenceMatcher(None, original_text, new_text)
        new_styles_map = []
        
        for style in styles_map:
            style_start = style['start']
            style_end = style['end']
            
            # Trouver les nouvelles positions pour ce style
            new_start = None
            new_end = None
            
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                # Si le style commence dans cette plage
                if i1 <= style_start < i2:
                    if tag == 'equal':
                        # Texte identique : position directe
                        new_start = j1 + (style_start - i1)
                    elif tag == 'replace':
                        # Texte remplacé : début de la nouvelle portion
                        new_start = j1
                    elif tag == 'delete':
                        # Texte supprimé : chercher la position suivante
                        new_start = j1
                    elif tag == 'insert':
                        # Insertion : ajuster
                        new_start = j1
                
                # Si le style se termine dans cette plage
                if i1 < style_end <= i2:
                    if tag == 'equal':
                        # Texte identique : position directe
                        new_end = j1 + (style_end - i1)
                    elif tag == 'replace':
                        # Texte remplacé : fin de la nouvelle portion
                        new_end = j2
                    elif tag == 'delete':
                        # Texte supprimé : position au début du bloc suivant
                        new_end = j1
                    elif tag == 'insert':
                        # Insertion : ajuster
                        new_end = j2
            
            # Si on a trouvé des positions valides
            if new_start is not None and new_end is not None and new_end > new_start:
                new_style = {
                    'start': new_start,
                    'end': min(new_end, len(new_text)),  # Ne pas dépasser la longueur du texte
                    'bold': style['bold'],
                    'italic': style['italic'],
                    'underline': style['underline'],
                    'font_name': style['font_name'],
                    'font_size': style['font_size'],
                    'font_color': style['font_color'],
                }
                new_styles_map.append(new_style)
        
        # Si aucun style n'a pu être mappé, utiliser le style dominant de l'original
        if not new_styles_map and styles_map:
            # Prendre le premier style (ou le dominant)
            dominant_style = styles_map[0]
            new_styles_map.append({
                'start': 0,
                'end': len(new_text),
                'bold': dominant_style['bold'],
                'italic': dominant_style['italic'],
                'underline': dominant_style['underline'],
                'font_name': dominant_style['font_name'],
                'font_size': dominant_style['font_size'],
                'font_color': dominant_style['font_color'],
            })
        
        return new_styles_map
    
    def _apply_styles_map(self, paragraph, new_text: str, styles_map: List[Dict]) -> None:
        """
        Applique une carte de styles à un paragraphe.
        
        Args:
            paragraph: Paragraphe docx
            new_text: Texte à insérer
            styles_map: Carte des styles à appliquer
        """
        # Supprimer tous les runs existants
        for run in paragraph.runs:
            run.text = ''
        
        if not styles_map:
            # Aucun style : créer un run simple
            paragraph.add_run(new_text)
            return
        
        # Trier les styles par position de début
        sorted_styles = sorted(styles_map, key=lambda x: x['start'])
        
        # Créer des runs pour chaque section de style
        last_end = 0
        
        for style in sorted_styles:
            start = max(style['start'], last_end)
            end = min(style['end'], len(new_text))
            
            if start >= len(new_text):
                break
            
            # Texte avant ce style (si gap)
            if start > last_end:
                gap_text = new_text[last_end:start]
                if gap_text:
                    paragraph.add_run(gap_text)
            
            # Texte avec ce style
            if end > start:
                styled_text = new_text[start:end]
                run = paragraph.add_run(styled_text)
                
                # Appliquer le style
                if style['bold'] is not None:
                    run.bold = style['bold']
                if style['italic'] is not None:
                    run.italic = style['italic']
                if style['underline'] is not None:
                    run.underline = style['underline']
                if style['font_name']:
                    run.font.name = style['font_name']
                if style['font_size']:
                    run.font.size = style['font_size']
                if style['font_color']:
                    run.font.color.rgb = style['font_color']
                
                last_end = end
        
        # Texte restant après tous les styles
        if last_end < len(new_text):
            remaining_text = new_text[last_end:]
            if remaining_text:
                run = paragraph.add_run(remaining_text)
                # Utiliser le style du dernier segment
                if sorted_styles:
                    last_style = sorted_styles[-1]
                    if last_style['bold'] is not None:
                        run.bold = last_style['bold']
                    if last_style['italic'] is not None:
                        run.italic = last_style['italic']
                    if last_style['underline'] is not None:
                        run.underline = last_style['underline']
                    if last_style['font_name']:
                        run.font.name = last_style['font_name']
                    if last_style['font_size']:
                        run.font.size = last_style['font_size']
                    if last_style['font_color']:
                        run.font.color.rgb = last_style['font_color']
    
    def _get_run_format(self, run) -> Dict:
        """
        Extrait le formatage d'un run.
        
        Args:
            run: Run docx
            
        Returns:
            Dictionnaire avec les propriétés de formatage
        """
        return {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color.rgb else None,
        }
    
    def _apply_run_format(self, run, format_dict: Dict) -> None:
        """
        Applique un formatage à un run.
        
        Args:
            run: Run docx
            format_dict: Dictionnaire de formatage
        """
        if format_dict['bold'] is not None:
            run.bold = format_dict['bold']
        if format_dict['italic'] is not None:
            run.italic = format_dict['italic']
        if format_dict['underline'] is not None:
            run.underline = format_dict['underline']
        if format_dict['font_name']:
            run.font.name = format_dict['font_name']
        if format_dict['font_size']:
            run.font.size = format_dict['font_size']
        if format_dict['font_color']:
            run.font.color.rgb = format_dict['font_color']
    
    def _get_dominant_format(self, paragraph) -> Dict:
        """
        Détermine le formatage dominant d'un paragraphe (pour le texte majoritaire).
        
        Args:
            paragraph: Paragraphe docx
            
        Returns:
            Dictionnaire de formatage dominant
        """
        # Compter la longueur de texte pour chaque format
        # Utilise une liste de dictionnaires au lieu d'un tuple comme clé
        format_list = []
        
        for run in paragraph.runs:
            if run.text.strip():  # Ignorer les runs vides
                # Stocker directement les objets (pas de conversion en string)
                run_format = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb if run.font.color.rgb else None,
                    'text_length': len(run.text)
                }
                format_list.append(run_format)
        
        # Trouver le format avec le plus de caractères
        if format_list:
            # Grouper par format similaire et sommer les longueurs
            format_groups = {}
            for fmt in format_list:
                # Créer une clé unique pour ce format
                key = (
                    fmt['bold'],
                    fmt['italic'],
                    fmt['underline'],
                    fmt['font_name'],
                    str(fmt['font_size']),
                    str(fmt['font_color'])
                )
                
                if key in format_groups:
                    format_groups[key]['total_length'] += fmt['text_length']
                else:
                    format_groups[key] = {
                        'bold': fmt['bold'],
                        'italic': fmt['italic'],
                        'underline': fmt['underline'],
                        'font_name': fmt['font_name'],
                        'font_size': fmt['font_size'],
                        'font_color': fmt['font_color'],
                        'total_length': fmt['text_length']
                    }
            
            # Trouver le groupe avec la plus grande longueur totale
            dominant = max(format_groups.values(), key=lambda x: x['total_length'])
            
            return {
                'bold': dominant['bold'],
                'italic': dominant['italic'],
                'underline': dominant['underline'],
                'font_name': dominant['font_name'],
                'font_size': dominant['font_size'],
                'font_color': dominant['font_color'],
            }
        
        # Format par défaut si aucun run avec texte
        return {
            'bold': False,
            'italic': False,
            'underline': False,
            'font_name': 'Calibri',
            'font_size': None,
            'font_color': None,
        }
    
    def _backup_paragraph_xml(self, paragraph):
        """
        Crée une sauvegarde XML du paragraphe pour restauration éventuelle.
        
        Args:
            paragraph: Paragraphe à sauvegarder
            
        Returns:
            Élément XML du paragraphe
        """
        from copy import deepcopy
        return deepcopy(paragraph._element)
    
    def _restore_paragraph_xml(self, paragraph, backup_xml):
        """
        Restaure un paragraphe depuis une sauvegarde XML.
        
        Args:
            paragraph: Paragraphe à restaurer
            backup_xml: Sauvegarde XML
        """
        paragraph._element.getparent().replace(paragraph._element, backup_xml)
        paragraph._element = backup_xml
    
    def _preserve_paragraph_format(self, original_paragraph, new_text: str) -> bool:
        """
        Remplace le texte d'un paragraphe en préservant tout le formatage.
        Utilise un système de mapping intelligent des styles caractère par caractère.
        
        Si le paragraphe contient des images, essaie de modifier et vérifie
        si les images sont toujours présentes. Si non, restaure le paragraphe original.
        
        Args:
            original_paragraph: Paragraphe docx original
            new_text: Nouveau texte à insérer
            
        Returns:
            True si la modification a été appliquée, False sinon
        """
        # Vérifier s'il y a des images
        has_images_before = self._has_images(original_paragraph)
        backup_xml = None
        
        if has_images_before:
            # Sauvegarder le paragraphe au cas où
            backup_xml = self._backup_paragraph_xml(original_paragraph)
            print("⚠️  IMAGES - Tentative de modification...", end=" ")
        
        # Sauvegarder les propriétés du paragraphe
        alignment = original_paragraph.alignment
        left_indent = original_paragraph.paragraph_format.left_indent
        right_indent = original_paragraph.paragraph_format.right_indent
        first_line_indent = original_paragraph.paragraph_format.first_line_indent
        space_before = original_paragraph.paragraph_format.space_before
        space_after = original_paragraph.paragraph_format.space_after
        line_spacing = original_paragraph.paragraph_format.line_spacing
        
        # NOUVEAU : Extraire la carte complète des styles avec positions
        original_text = original_paragraph.text
        styles_map = self._extract_styles_map(original_paragraph)
        
        # NOUVEAU : Mapper les styles sur le nouveau texte en utilisant difflib
        new_styles_map = self._map_styles_to_new_text(original_text, new_text, styles_map)
        
        # NOUVEAU : Appliquer les styles mappés
        self._apply_styles_map(original_paragraph, new_text, new_styles_map)
        
        # Restaurer les propriétés du paragraphe
        original_paragraph.alignment = alignment
        if left_indent is not None:
            original_paragraph.paragraph_format.left_indent = left_indent
        if right_indent is not None:
            original_paragraph.paragraph_format.right_indent = right_indent
        if first_line_indent is not None:
            original_paragraph.paragraph_format.first_line_indent = first_line_indent
        if space_before is not None:
            original_paragraph.paragraph_format.space_before = space_before
        if space_after is not None:
            original_paragraph.paragraph_format.space_after = space_after
        if line_spacing is not None:
            original_paragraph.paragraph_format.line_spacing = line_spacing
        
        # Si il y avait des images, vérifier qu'elles sont toujours là
        if has_images_before:
            has_images_after = self._has_images(original_paragraph)
            
            if not has_images_after:
                # Les images ont disparu ! Restaurer le paragraphe original
                print("❌ Images perdues, RESTAURATION !", end=" ")
                self._restore_paragraph_xml(original_paragraph, backup_xml)
                return False
            else:
                # Les images sont toujours là !
                print("✅ Images préservées !", end=" ")
                return True
        
        return True
    
    def _verify_images(self) -> None:
        """
        Vérifie que toutes les images sont toujours présentes après le traitement.
        """
        current_image_count = 0
        current_paragraphs_with_images = []
        
        for i, paragraph in enumerate(self.current_document.paragraphs):
            if self._has_images(paragraph):
                current_paragraphs_with_images.append(i + 1)
                for run in paragraph.runs:
                    if hasattr(run, '_element'):
                        for child in run._element:
                            if 'drawing' in child.tag or 'pict' in child.tag:
                                current_image_count += 1
        
        print("\n" + "=" * 60)
        print("VÉRIFICATION DES IMAGES")
        print("=" * 60)
        print(f"Images au début: {self.initial_image_count}")
        print(f"Images maintenant: {current_image_count}")
        
        if current_image_count == self.initial_image_count:
            print("✅ TOUTES LES IMAGES SONT PRÉSERVÉES !")
        else:
            print(f"❌ ATTENTION: {self.initial_image_count - current_image_count} image(s) perdue(s) !")
            print(f"   Paragraphes avec images au début: {self.paragraphs_with_images}")
            print(f"   Paragraphes avec images maintenant: {current_paragraphs_with_images}")
        
        if len(self.paragraphs_with_images) > 0:
            print(f"\nℹ️  {len(self.paragraphs_with_images)} paragraphe(s) avec images n'ont PAS été modifiés:")
            print(f"   Paragraphes: {', '.join(map(str, self.paragraphs_with_images))}")
        
        print("=" * 60)
    
    def _detect_language(self) -> None:
        """
        Détecte la langue du document en analysant un échantillon de texte.
        """
        try:
            # Concaténer plusieurs paragraphes pour avoir un échantillon représentatif
            sample_text = " ".join([
                p.strip() for p in self.paragraphs_cache 
                if p.strip() and len(p.strip()) > 20
            ][:10])  # Prendre les 10 premiers paragraphes significatifs
            
            if sample_text:
                self.detected_language = detect(sample_text)
        except Exception as e:
            print(f"  ⚠️ Détection de langue échouée: {e}")
            self.detected_language = None
    
    def _get_language_name(self, lang_code: str) -> str:
        """
        Convertit un code de langue en nom complet.
        
        Args:
            lang_code: Code de langue (ex: 'fr', 'en', 'es')
            
        Returns:
            Nom de la langue
        """
        language_map = {
            'fr': 'Français',
            'en': 'Anglais',
            'es': 'Espagnol',
            'de': 'Allemand',
            'it': 'Italien',
            'pt': 'Portugais',
            'nl': 'Néerlandais',
            'ru': 'Russe',
            'zh-cn': 'Chinois (simplifié)',
            'zh-tw': 'Chinois (traditionnel)',
            'ja': 'Japonais',
            'ko': 'Coréen',
            'ar': 'Arabe',
            'tr': 'Turc',
            'pl': 'Polonais',
            'sv': 'Suédois',
            'da': 'Danois',
            'no': 'Norvégien',
            'fi': 'Finnois',
        }
        return language_map.get(lang_code, f"Langue inconnue ({lang_code})")
    
    def _init_log_file(self) -> None:
        """
        Initialise le fichier de log pour ce document.
        Crée le dossier LOGS s'il n'existe pas.
        """
        # Créer le dossier LOGS
        log_dir = Path("LOGS")
        log_dir.mkdir(exist_ok=True)
        
        # Nom du fichier : nom_document_YYYYMMDD.txt
        date_str = datetime.now().strftime("%Y%m%d")
        doc_name = self.current_path.stem if self.current_path else "document"
        log_filename = f"{doc_name}_{date_str}.txt"
        
        self.log_file = log_dir / log_filename
        
        # Créer/initialiser le fichier de log
        with open(self.log_file, 'a', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write(f"LOG DE MODIFICATIONS - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Document: {self.current_path.name if self.current_path else 'inconnu'}\n")
            f.write(f"Nombre de paragraphes: {len(self.paragraphs_cache)}\n")
            if self.detected_language:
                f.write(f"Langue détectée: {self._get_language_name(self.detected_language)}\n")
            f.write("=" * 80 + "\n\n")
        
        print(f"  Log initialisé: {self.log_file}")
    
    def _detect_differences(self, original: str, modified: str) -> List[Dict[str, any]]:
        """
        Détecte les différences entre deux textes en utilisant difflib.
        
        Args:
            original: Texte original
            modified: Texte modifié
            
        Returns:
            Liste de dictionnaires décrivant chaque différence
        """
        differences = []
        
        # Utiliser SequenceMatcher pour détecter les changements
        matcher = difflib.SequenceMatcher(None, original, modified)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                differences.append({
                    'type': 'REMPLACEMENT',
                    'position': i1,
                    'original': original[i1:i2],
                    'modifie': modified[j1:j2],
                    'contexte_avant': original[max(0, i1-20):i1],
                    'contexte_apres': original[i2:min(len(original), i2+20)]
                })
            elif tag == 'delete':
                differences.append({
                    'type': 'SUPPRESSION',
                    'position': i1,
                    'original': original[i1:i2],
                    'contexte_avant': original[max(0, i1-20):i1],
                    'contexte_apres': original[i2:min(len(original), i2+20)]
                })
            elif tag == 'insert':
                differences.append({
                    'type': 'AJOUT',
                    'position': i1,
                    'modifie': modified[j1:j2],
                    'contexte_avant': original[max(0, i1-20):i1] if i1 > 0 else '',
                    'contexte_apres': original[i1:min(len(original), i1+20)] if i1 < len(original) else ''
                })
        
        return differences
    
    def _log_change(self, paragraph_num: int, original: str, modified: str, instruction: str) -> None:
        """
        Enregistre un changement dans le fichier de log.
        
        Args:
            paragraph_num: Numéro du paragraphe
            original: Texte original
            modified: Texte modifié
            instruction: Instruction qui a causé le changement
        """
        if not self.log_file:
            return
        
        # Déterminer si c'est une correction pour analyser les différences
        is_correction = any(word in instruction.lower() for word in ['corrige', 'correction', 'orthographe', 'grammaire'])
        
        with open(self.log_file, 'a', encoding='utf-8') as f:
            f.write("-" * 80 + "\n")
            f.write(f"PARAGRAPHE {paragraph_num}\n")
            f.write(f"Instruction: {instruction}\n")
            f.write(f"Date/Heure: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("-" * 80 + "\n\n")
            
            # Si c'est une correction, détecter et afficher les différences
            if is_correction and original != modified:
                differences = self._detect_differences(original, modified)
                
                if differences:
                    f.write(f"NOMBRE DE MODIFICATIONS: {len(differences)}\n\n")
                    
                    for i, diff in enumerate(differences, 1):
                        f.write(f"  [{i}] {diff['type']}\n")
                        f.write(f"      Position: caractère {diff['position']}\n")
                        
                        if 'contexte_avant' in diff and diff['contexte_avant']:
                            f.write(f"      Contexte avant: ...{diff['contexte_avant']}\n")
                        
                        if diff['type'] == 'REMPLACEMENT':
                            f.write(f"      AVANT: '{diff['original']}'\n")
                            f.write(f"      APRES: '{diff['modifie']}'\n")
                        elif diff['type'] == 'SUPPRESSION':
                            f.write(f"      SUPPRIME: '{diff['original']}'\n")
                        elif diff['type'] == 'AJOUT':
                            f.write(f"      AJOUTE: '{diff['modifie']}'\n")
                        
                        if 'contexte_apres' in diff and diff['contexte_apres']:
                            f.write(f"      Contexte après: {diff['contexte_apres']}...\n")
                        
                        f.write("\n")
                else:
                    f.write("AUCUNE DIFFÉRENCE DÉTECTÉE (textes identiques)\n\n")
            
            # Toujours afficher le avant/après complet
            f.write("TEXTE ORIGINAL:\n")
            f.write("-" * 40 + "\n")
            f.write(original + "\n")
            f.write("-" * 40 + "\n\n")
            
            f.write("TEXTE MODIFIE:\n")
            f.write("-" * 40 + "\n")
            f.write(modified + "\n")
            f.write("-" * 40 + "\n\n")
            
            f.write("\n")
    
    def _call_openai(self, instruction: str, text: str, context: str = "", is_correction: bool = False) -> str:
        """
        Appelle l'API OpenAI pour traiter du texte.
        
        Args:
            instruction: Instruction à donner au modèle
            text: Texte à traiter
            context: Contexte additionnel
            is_correction: Si True, ajoute la langue détectée au contexte
            
        Returns:
            Texte traité
        """
        system_content = (
            "Tu es un assistant expert en révision de documents. "
            "Tu dois UNIQUEMENT retourner le texte modifié, sans explications, "
            "sans commentaires, sans formatage markdown. "
            "Préserve la structure exacte du texte (sauts de ligne, espaces, etc.)."
        )
        
        # Ajouter la langue au contexte système si c'est une correction
        if is_correction and self.detected_language:
            language_name = self._get_language_name(self.detected_language)
            system_content += f"\nLe document est en {language_name}. Effectue la correction dans cette langue."
        
        messages = [
            {
                "role": "system",
                "content": system_content
            }
        ]
        
        # Ajouter l'historique de conversation (limité aux 5 derniers messages)
        if self.conversation_history:
            messages.extend(self.conversation_history[-5:])
        
        # Ajouter le contexte si fourni
        if context:
            messages.append({
                "role": "system",
                "content": f"Contexte: {context}"
            })
        
        # Ajouter l'instruction actuelle
        messages.append({
            "role": "user",
            "content": f"{instruction}\n\nTexte:\n{text}"
        })
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.3,
            )
            
            result = response.choices[0].message.content.strip()
            
            # Ajouter à l'historique
            self.conversation_history.append({
                "role": "user",
                "content": f"{instruction} (paragraphe)"
            })
            self.conversation_history.append({
                "role": "assistant",
                "content": result[:100] + "..." if len(result) > 100 else result
            })
            
            return result
            
        except Exception as e:
            print(f"❌ Erreur API OpenAI: {e}")
            return text  # Retourner le texte original en cas d'erreur
    
    def process_document(self, instruction: str, batch_size: int = 5) -> None:
        """
        Traite le document avec l'instruction donnée.
        
        Args:
            instruction: Instruction à exécuter (ex: "corrige les fautes", "traduis en anglais")
            batch_size: Nombre de paragraphes à traiter ensemble pour le contexte
        """
        if not self.current_document:
            raise ValueError("Aucun document chargé. Utilisez load_document() d'abord.")
        
        # Détecter si c'est une correction pour ajouter la langue au contexte
        is_correction = any(word in instruction.lower() for word in ['corrige', 'correction', 'orthographe', 'grammaire'])
        
        print(f"\n🔄 Traitement: {instruction}")
        if is_correction and self.detected_language:
            print(f"   Langue: {self._get_language_name(self.detected_language)}")
        print("=" * 60)
        
        paragraphs = list(self.current_document.paragraphs)
        total = len(paragraphs)
        
        for i, paragraph in enumerate(paragraphs):
            if not paragraph.text.strip():
                continue  # Ignorer les paragraphes vides
            
            # Créer un contexte avec les paragraphes précédents
            context_start = max(0, i - 2)
            context_paragraphs = [p.text for p in paragraphs[context_start:i] if p.text.strip()]
            context = " [...] ".join(context_paragraphs[-2:]) if context_paragraphs else ""
            
            print(f"Paragraphe {i+1}/{total}...", end=" ")
            
            original_text = paragraph.text
            processed_text = self._call_openai(instruction, original_text, context, is_correction)
            
            # Mettre à jour le paragraphe en préservant le format
            if processed_text and processed_text != original_text:
                # _preserve_paragraph_format retourne True si la modification a été appliquée
                modification_applied = self._preserve_paragraph_format(paragraph, processed_text)
                
                if modification_applied:
                    # Logger le changement seulement si la modification a été appliquée
                    self._log_change(i + 1, original_text, processed_text, instruction)
                    print("✓ Modifié")
                else:
                    print("○ Non modifié (images)")
            else:
                print("○ Inchangé")
        
        print("=" * 60)
        print("✓ Traitement terminé !")
        
        # Vérifier les images après traitement
        self._verify_images()
    
    def save_document(self, output_path: Optional[str] = None) -> None:
        """
        Sauvegarde le document modifié.
        
        Args:
            output_path: Chemin de sortie (si None, ajoute "_modifié" au nom original)
        """
        if not self.current_document:
            raise ValueError("Aucun document à sauvegarder.")
        
        if output_path is None:
            output_path = self.current_path.parent / f"{self.current_path.stem}_modifié{self.current_path.suffix}"
        
        self.current_document.save(output_path)
        print(f"\n💾 Document sauvegardé: {output_path}")
    
    def interactive_mode(self, file_path: str) -> None:
        """
        Mode interactif pour traiter un document avec plusieurs commandes.
        
        Args:
            file_path: Chemin vers le document à traiter
        """
        self.load_document(file_path)
        
        print("\n" + "=" * 60)
        print("MODE INTERACTIF - Document Reviewer")
        print("=" * 60)
        print("\nCommandes disponibles:")
        print("  - 'corrige' : Corrige les fautes d'orthographe")
        print("  - 'traduis [langue]' : Traduit le document")
        print("  - 'améliore' : Améliore le style")
        print("  - 'résume' : Résume le contenu")
        print("  - ou toute autre instruction personnalisée")
        print("  - 'save' : Sauvegarder les modifications")
        print("  - 'quit' : Quitter sans sauvegarder")
        print("  - 'save+quit' : Sauvegarder et quitter")
        print("=" * 60)
        
        while True:
            try:
                user_input = input("\n➤ Votre commande: ").strip()
                
                if not user_input:
                    continue
                
                if user_input.lower() == 'quit':
                    print("Au revoir !")
                    break
                
                if user_input.lower() == 'save':
                    self.save_document()
                    continue
                
                if user_input.lower() == 'save+quit':
                    self.save_document()
                    print("Au revoir !")
                    break
                
                # Traiter l'instruction
                if user_input.lower().startswith('corrige'):
                    instruction = "Corrige toutes les fautes d'orthographe et de grammaire dans ce texte."
                elif user_input.lower().startswith('traduis'):
                    langue = user_input.split(maxsplit=1)[1] if ' ' in user_input else "anglais"
                    instruction = f"Traduis ce texte en {langue}."
                elif user_input.lower() == 'améliore':
                    instruction = "Améliore le style et la clarté de ce texte."
                elif user_input.lower() == 'résume':
                    instruction = "Résume ce texte de manière concise."
                else:
                    instruction = user_input
                
                self.process_document(instruction)
                
            except KeyboardInterrupt:
                print("\n\nInterruption détectée.")
                save = input("Voulez-vous sauvegarder avant de quitter ? (o/n): ").strip().lower()
                if save == 'o':
                    self.save_document()
                break
            except Exception as e:
                print(f"❌ Erreur: {e}")


def main():
    """Fonction principale."""
    print("=" * 60)
    print("DOCUMENT REVIEWER - Correction avec OpenAI")
    print("=" * 60)
    
    # Demander la clé API si non définie
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("\n⚠️  Clé API OpenAI non trouvée.")
        print("Options:")
        print("  1. Créez un fichier .env avec: OPENAI_API_KEY=votre-clé")
        print("  2. Définissez la variable d'environnement OPENAI_API_KEY")
        print("  3. Entrez-la maintenant:")
        api_key = input("➤ Clé API: ").strip()
        
        if not api_key:
            print("❌ Clé API requise pour fonctionner.")
            return
    else:
        print("✓ Clé API OpenAI chargée depuis l'environnement")
    
    # Créer le reviewer
    reviewer = DocumentReviewer(api_key=api_key)
    
    # Demander le fichier
    file_path = input("\n➤ Chemin du document Word: ").strip().strip('"')
    
    if not file_path:
        print("❌ Aucun fichier spécifié.")
        return
    
    # Lancer le mode interactif
    try:
        reviewer.interactive_mode(file_path)
    except Exception as e:
        print(f"❌ Erreur: {e}")


if __name__ == "__main__":
    main()

