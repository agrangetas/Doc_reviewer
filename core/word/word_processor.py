"""
Word Processor - Implémentation complète
Traitement de documents Word avec IA.
"""

from pathlib import Path
from typing import Optional
from docx import Document
from datetime import datetime

from core.base.document_processor import DocumentProcessor


class WordProcessor(DocumentProcessor):
    """Processeur pour les documents Word (.docx)."""
    
    def __init__(self, config, image_handler, style_extractor, style_mapper,
                 language_detector, ai_processor, logger, style_uniformizer):
        """Initialise le processeur Word."""
        super().__init__(
            config, image_handler, style_extractor, style_mapper,
            language_detector, ai_processor, logger, style_uniformizer
        )
        self.initial_image_count = 0
        self.paragraphs_with_images = []
    
    def load_document(self, file_path: str) -> None:
        """
        Charge un document Word.
        
        Args:
            file_path: Chemin vers le fichier .docx
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Le fichier {file_path} n'existe pas.")
        
        self.current_document = Document(file_path)
        self.current_path = path
        
        # Détecter la langue
        sample_text = " ".join([p.text.strip() for p in self.current_document.paragraphs if p.text.strip()][:10])
        if sample_text:
            lang_code = self.language_detector.detect_language(sample_text)
            if lang_code:
                self.detected_language = lang_code
        
        # Compter les images
        self.initial_image_count, self.paragraphs_with_images = self.image_handler.count_images(self.current_document)
        
        # Afficher les infos
        print(f"✓ Document chargé: {path.name}")
        print(f"  Nombre de paragraphes: {len(self.current_document.paragraphs)}")
        print(f"  Modèle OpenAI: {self.ai_processor.model}")
        if self.detected_language:
            print(f"  Langue détectée: {self.language_detector.get_language_name(self.detected_language)}")
        
        # Initialiser le log
        doc_info = {
            'paragraph_count': len(self.current_document.paragraphs),
            'language': self.language_detector.get_language_name(self.detected_language) if self.detected_language else None
        }
        self.logger.init_log_file(path.name, doc_info)
        
        # Afficher info images
        if self.initial_image_count > 0:
            print(f"  Images trouvées: {self.initial_image_count} image(s) dans {len(self.paragraphs_with_images)} paragraphe(s)")
            print(f"  ⚠️  Les paragraphes avec images seront traités avec précaution")
    
    def save_document(self, output_path: Optional[str] = None) -> None:
        """
        Sauvegarde le document Word.
        
        Args:
            output_path: Chemin de sortie (optionnel)
        """
        if not self.current_document:
            raise ValueError("Aucun document à sauvegarder.")
        
        if output_path is None:
            output_path = self.current_path.parent / f"{self.current_path.stem}_modifié{self.current_path.suffix}"
        
        self.current_document.save(output_path)
        print(f"\n💾 Document sauvegardé: {output_path}")
    
    def process_document(self, instruction: str) -> None:
        """
        Traite le document avec une instruction.
        
        Args:
            instruction: Instruction à exécuter
        """
        if not self.current_document:
            raise ValueError("Aucun document chargé.")
        
        is_correction = any(word in instruction.lower() for word in ['corrige', 'correction', 'orthographe', 'grammaire'])
        
        from features.language_detector import LanguageDetector
        language_name = LanguageDetector.get_language_name(self.detected_language) if self.detected_language else None
        
        print(f"\n🔄 Traitement: {instruction}")
        if is_correction and language_name:
            print(f"   Langue: {language_name}")
        print("=" * 60)
        
        modified_count = 0
        paragraphs = self.current_document.paragraphs
        total_paragraphs = len(paragraphs)
        
        for i, paragraph in enumerate(paragraphs, 1):
            if not paragraph.text.strip():
                continue
            
            print(f"Paragraphe {i}/{total_paragraphs}...", end=" ")
            
            try:
                # Contexte : paragraphes précédents
                context = " [...] ".join([
                    p.text.strip() 
                    for p in paragraphs[max(0, i-3):i-1] 
                    if p.text.strip()
                ])
                
                # Si images détectées
                has_images = i in self.paragraphs_with_images
                if has_images:
                    # Backup
                    backup = self.image_handler.backup_paragraph(paragraph)
                
                original_text = paragraph.text
                processed_text = self.ai_processor.call_openai(
                    instruction, original_text, context, is_correction, language_name
                )
                
                if processed_text and processed_text != original_text:
                    # Extraire les styles
                    styles_map = self.style_extractor.extract_styles_map(paragraph)
                    new_styles_map = self.style_mapper.map_styles_to_new_text(
                        original_text, processed_text, styles_map
                    )
                    
                    # Appliquer
                    modification_applied = self._preserve_paragraph_format(
                        paragraph, processed_text, new_styles_map
                    )
                    
                    # Vérifier les images
                    if has_images:
                        images_after = self.image_handler.count_images_in_paragraph(paragraph)
                        images_before = self.image_handler.count_images_in_paragraph_from_backup(backup)
                        
                        if images_after != images_before:
                            # Restaurer
                            self.image_handler.restore_paragraph(paragraph, backup)
                            print("○ Non modifié (images)")
                            modification_applied = False
                    
                    if modification_applied:
                        self.logger.log_change(i, original_text, processed_text, instruction)
                        print("✓ Modifié")
                        modified_count += 1
                else:
                    print("○ Inchangé")
            
            except Exception as e:
                print(f"❌ Erreur: {e}")
        
        # Vérification finale des images
        if self.initial_image_count > 0:
            final_count, _ = self.image_handler.count_images(self.current_document)
            if final_count != self.initial_image_count:
                print(f"\n⚠️  ATTENTION: {self.initial_image_count - final_count} image(s) manquante(s) !")
            else:
                print(f"\n✓ Toutes les images préservées ({final_count})")
        
        print("=" * 60)
        print(f"✓ Traitement terminé ! ({modified_count} paragraphes modifiés)")
    
    def _preserve_paragraph_format(self, paragraph, new_text: str, styles_map: list) -> bool:
        """
        Applique le nouveau texte en préservant le format.
        
        Args:
            paragraph: Paragraphe à modifier
            new_text: Nouveau texte
            styles_map: Carte des styles
            
        Returns:
            True si modifié, False sinon
        """
        try:
            self.style_mapper.apply_styles_map(paragraph, new_text, styles_map)
            return True
        except Exception as e:
            print(f"⚠️  Erreur appliation styles: {e}")
            return False
    
    def process_targeted(self, target, instruction: str) -> None:
        """
        Traite un élément ciblé spécifiquement.
        
        Args:
            target: ResolvedTarget avec le paragraphe ciblé
            instruction: Instruction à appliquer
        """
        if not self.current_document:
            raise ValueError("Aucun document chargé.")
        
        paragraph_num = target.paragraph
        if not paragraph_num:
            raise ValueError("Aucun paragraphe ciblé.")
        
        # Vérifier que le paragraphe existe
        if paragraph_num < 1 or paragraph_num > len(self.current_document.paragraphs):
            raise ValueError(f"Paragraphe {paragraph_num} n'existe pas (document a {len(self.current_document.paragraphs)} paragraphes).")
        
        paragraph = self.current_document.paragraphs[paragraph_num - 1]  # Index 0-based
        
        if not paragraph.text.strip():
            print(f"⚠️  Paragraphe {paragraph_num} est vide, ignoré.")
            return
        
        print(f"\n🎯 Traitement ciblé: Paragraphe {paragraph_num}")
        print(f"   Instruction: {instruction}")
        print("=" * 60)
        
        from features.language_detector import LanguageDetector
        is_correction = any(word in instruction.lower() for word in ['corrige', 'correction', 'orthographe', 'grammaire'])
        language_name = LanguageDetector.get_language_name(self.detected_language) if self.detected_language else None
        
        # Contexte : paragraphes voisins
        context_parts = []
        for i in range(max(0, paragraph_num - 3), paragraph_num - 1):
            p = self.current_document.paragraphs[i]
            if p.text.strip():
                context_parts.append(p.text.strip())
        context = " [...] ".join(context_parts)
        
        # Vérifier images
        has_images = paragraph_num in self.paragraphs_with_images
        if has_images:
            backup = self.image_handler.backup_paragraph(paragraph)
        
        try:
            original_text = paragraph.text
            print(f"Texte original: {original_text[:100]}{'...' if len(original_text) > 100 else ''}")
            
            # Traitement
            processed_text = self.ai_processor.call_openai(
                instruction, original_text, context, is_correction, language_name
            )
            
            if processed_text and processed_text != original_text:
                # Extraire et mapper les styles
                styles_map = self.style_extractor.extract_styles_map(paragraph)
                new_styles_map = self.style_mapper.map_styles_to_new_text(
                    original_text, processed_text, styles_map
                )
                
                # Appliquer
                modification_applied = self._preserve_paragraph_format(
                    paragraph, processed_text, new_styles_map
                )
                
                # Vérifier les images
                if has_images:
                    images_after = self.image_handler.count_images_in_paragraph(paragraph)
                    images_before = self.image_handler.count_images_in_paragraph_from_backup(backup)
                    
                    if images_after != images_before:
                        # Restaurer
                        self.image_handler.restore_paragraph(paragraph, backup)
                        print("○ Non modifié (images préservées)")
                        modification_applied = False
                
                if modification_applied:
                    self.logger.log_change(paragraph_num, original_text, processed_text, 
                                          f"{instruction} (ciblé)")
                    print(f"✓ Paragraphe {paragraph_num} modifié")
                    print(f"Nouveau texte: {processed_text[:100]}{'...' if len(processed_text) > 100 else ''}")
            else:
                print("○ Aucun changement")
        
        except Exception as e:
            print(f"❌ Erreur: {e}")
        
        print("=" * 60)
    
    def uniformize_styles(self) -> None:
        """Uniformise les styles du document Word."""
        if not self.current_document:
            raise ValueError("Aucun document chargé.")
        
        result = self.style_uniformizer.uniformize(self.current_document)
        
        # Logger l'opération
        if self.logger.log_file and not result.get('cancelled'):
            with open(self.logger.log_file, 'a', encoding='utf-8') as f:
                f.write("-" * 80 + "\n")
                f.write(f"UNIFORMISATION DES STYLES\n")
                f.write(f"Date/Heure: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("-" * 80 + "\n\n")
                
                # Styles appliqués
                target_font = result.get('target_font', 'N/A')
                target_size_pt = result.get('target_size_pt', 'N/A')
                target_color = result.get('target_color', 'Inchangée')
                target_spacing = result.get('target_line_spacing', 'Inchangée')
                
                f.write(f"Police cible: {target_font}\n")
                f.write(f"Taille cible: {target_size_pt}\n")
                f.write(f"Couleur cible (texte): {target_color}\n")
                f.write(f"Interligne cible: {target_spacing}\n")
                f.write(f"\nModifications appliquées:\n")
                f.write(f"  Paragraphes modifiés: {result.get('modified_paragraphs', 0)}\n")
                f.write(f"  Polices changées: {result.get('font_changes', 0)}\n")
                f.write(f"  Tailles changées: {result.get('size_changes', 0)}\n")
                f.write(f"  Couleurs changées: {result.get('color_changes', 0)}\n")
                f.write(f"  Interlignes changés: {result.get('spacing_changes', 0)}\n")
                f.write("\n" + "=" * 80 + "\n\n")
    
    def get_format_name(self) -> str:
        """Retourne 'Word'."""
        return "Word"

